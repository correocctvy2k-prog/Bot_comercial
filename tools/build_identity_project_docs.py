from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path
from datetime import date

OUT = Path("entregables/proyecto_identidad")
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "16324F"; BLUE = "2E74B5"; LIGHT = "EAF1F7"; PALE = "F4F6F9"
GRAY = "5B6573"; DARK = "17202A"; GREEN = "1F6B52"; GOLD = "8A6500"; RED = "9B1C1C"

SOURCES = [
    ("Microsoft: retiro de SMS/voz y adopción de passkeys", "https://learn.microsoft.com/en-us/entra/identity/authentication/concept-sms-voice-retirement"),
    ("Microsoft: passkeys (FIDO2) en Entra ID", "https://learn.microsoft.com/en-us/entra/identity/authentication/concept-authentication-passkeys-fido2"),
    ("Microsoft: sincronización de hash de contraseña", "https://learn.microsoft.com/en-us/entra/identity/hybrid/connect/how-to-connect-password-hash-synchronization"),
    ("Microsoft: Windows Hello for Business y Cloud Kerberos Trust", "https://learn.microsoft.com/en-us/windows/security/identity-protection/hello-for-business/deploy/"),
    ("Microsoft: comprobación de Microsoft Entra Hybrid Join", "https://learn.microsoft.com/es-es/entra/identity/devices/how-to-hybrid-join-verify"),
    ("Microsoft: recomendaciones de políticas de contraseña", "https://learn.microsoft.com/en-us/microsoft-365/admin/misc/password-policy-recommendations?view=o365-worldwide"),
    ("Microsoft Research: Microsoft Password Guidance", "https://www.microsoft.com/en-us/research/wp-content/uploads/2016/06/Microsoft_Password_Guidance-1.pdf"),
    ("NIST SP 800-63B-4: Authentication and Authenticator Management", "https://nvlpubs.nist.gov/nistpubs/SpecialPublications/NIST.SP.800-63b-4.pdf"),
    ("Fortinet: SSL VPN con Microsoft Entra SSO", "https://docs.fortinet.com/document/fortigate/7.4.8/administration-guide/109964/ssl-vpn-with-microsoft-entra-sso-integration"),
    ("Fortinet: VPN antes del inicio de sesión", "https://docs.fortinet.com/document/forticlient/7.4.3/administration-guide/479513/activating-vpn-before-windows-logon"),
    ("Fortinet: prelogon con certificado de equipo y EMS", "https://docs.fortinet.com/document/forticlient/7.4.4/administration-guide/505235/appendix-f-ssl-vpn-prelogon"),
    ("Microsoft: precios de Microsoft Entra", "https://www.microsoft.com/en-us/security/business/microsoft-entra-pricing"),
    ("Microsoft: precio y contenido de Business Premium", "https://www.microsoft.com/en-us/microsoft-365/business/microsoft-365-business-premium"),
]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for tag, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+tag))
        if node is None: node = OxmlElement('w:'+tag); tcMar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'), 'dxa')

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def set_cell_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'),'dxa')

def set_table_geometry(t, widths, indent=120):
    tblPr=t._tbl.tblPr
    tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    tblInd=tblPr.find(qn('w:tblInd'))
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),str(indent)); tblInd.set(qn('w:type'),'dxa')
    grid=t._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(width)); grid.append(col)
    for row in t.rows:
        for idx,cell in enumerate(row.cells): set_cell_width(cell,widths[idx])

def hyperlink(paragraph, text, url):
    part = paragraph.part; rid = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    h = OxmlElement('w:hyperlink'); h.set(qn('r:id'), rid)
    r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), BLUE); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    r.append(rPr); t = OxmlElement('w:t'); t.text = text; r.append(t); h.append(r); paragraph._p.append(h)

def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Página '); run.font.size = Pt(9); run.font.color.rgb = RGBColor.from_string(GRAY)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); paragraph._p.append(fld)

def setup(doc, label):
    sec = doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(DARK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
    for name,size,color,before,after in [('Title',26,NAVY,0,8),('Subtitle',14,GRAY,0,18),('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,6),('Heading 3',12,NAVY,8,4)]:
        s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=name!='Subtitle'
        s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for st in ['List Bullet','List Number']:
        styles[st].font.name='Calibri'; styles[st].font.size=Pt(11); styles[st].paragraph_format.space_after=Pt(5); styles[st].paragraph_format.line_spacing=1.10
    header=sec.header.paragraphs[0]; header.text=label; header.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in header.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GRAY)
    page_number(sec.footer.paragraphs[0])
    return doc

def title(doc, name, subtitle, doc_type):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(30); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(doc_type.upper()); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(BLUE)
    doc.add_paragraph(name, 'Title'); doc.add_paragraph(subtitle, 'Subtitle')
    p=doc.add_paragraph(); p.add_run('Preparado para: ').bold=True; p.add_run('Dirección y equipo de Tecnología')
    p=doc.add_paragraph(); p.add_run('Fecha base: ').bold=True; p.add_run('11 de agosto de 2026')
    p=doc.add_paragraph(); p.add_run('Estado: ').bold=True; p.add_run('Propuesta para aprobación y levantamiento técnico')
    doc.add_paragraph('Documento sujeto a validación de licencias Microsoft, versión de FortiOS/FortiClient, contrato Fortinet y capacidad interna.', 'Caption')

def callout(doc, heading, text, fill=LIGHT, color=NAVY):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    set_repeat_header(t.rows[0])
    c=t.cell(0,0); set_cell_width(c,9360); margins(c,150,180,150,180); shade(c,fill)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3); r=p.add_run(heading); r.bold=True; r.font.color.rgb=RGBColor.from_string(color)
    p=c.add_paragraph(text); p.paragraph_format.space_after=Pt(0)
    set_table_geometry(t,[9360])
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def table(doc, headers, rows, widths=None, font=9.5):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    t.style='Table Grid'; set_repeat_header(t.rows[0])
    if widths is None: widths=[9360//len(headers)]*len(headers)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_cell_width(c,widths[i]); margins(c); shade(c,LIGHT); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(h); r.bold=True; r.font.size=Pt(font); r.font.color.rgb=RGBColor.from_string(NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            set_cell_width(cells[i],widths[i]); margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(val)); r.font.size=Pt(font)
    set_table_geometry(t,widths)
    return t

def bullets(doc, items):
    for x in items: doc.add_paragraph(x, 'List Bullet')

def numbered(doc, items):
    for x in items: doc.add_paragraph(x, 'List Number')

def sources(doc):
    doc.add_heading('Fuentes y respaldo', level=1)
    doc.add_paragraph('Fuentes primarias consultadas. Los precios son referencias públicas en USD observadas el 11 de agosto de 2026; deben confirmarse con el proveedor colombiano antes de aprobar compras.')
    for label,url in SOURCES:
        p=doc.add_paragraph(style='List Bullet'); hyperlink(p,label,url)

def executive():
    doc=setup(Document(),'Proyecto de identidad y acceso seguro | Documento ejecutivo')
    title(doc,'Modernización de identidad y acceso seguro','Caso ejecutivo para decisión directiva','Documento ejecutivo')
    callout(doc,'Decisión solicitada','Autorizar un proyecto por fases para unificar la identidad de cada empleado, adoptar autenticación resistente al phishing, extender los controles corporativos a portátiles y modernizar la política de contraseñas sin reemplazar Active Directory, Zimbra ni FortiGate.')
    doc.add_heading('1. Resumen ejecutivo',1)
    doc.add_paragraph('La organización dispone de una base local sólida: tres controladores de dominio Windows Server 2022 para aproximadamente 40 usuarios, correo independiente en Zimbra, Microsoft 365 para una parte de la población y acceso remoto mediante FortiClient VPN seguido de Escritorio remoto. El riesgo principal no es la disponibilidad del dominio, sino la fragmentación de identidad y la dependencia de credenciales susceptibles de phishing.')
    bullets(doc,[
        'Una misma persona utiliza identidades diferentes en Windows, Microsoft 365, Zimbra y el portátil remoto.',
        'Las licencias de Microsoft 365 se administran con direcciones basadas en cargos; esto reduce trazabilidad y puede inducir uso compartido de credenciales.',
        'Los portátiles externos se inician con cuentas locales antes de establecer la VPN, por lo que quedan parcialmente fuera del gobierno corporativo.',
        'SMS y llamada dejarán de ser métodos provistos por Microsoft Entra el 1 de febrero de 2027.',
        'El cambio periódico obligatorio de contraseña incentiva variaciones predecibles y no debe ser el control central de seguridad.'
    ])
    doc.add_heading('2. Estado actual',1)
    table(doc,['Componente','Situación actual','Consecuencia'],[
        ('Active Directory','MIDOMINIO.LOCAL; tres DC Server 2022; ~40 usuarios.','Base estable, pero separada del tenant.'),
        ('Microsoft 365','Tenant independiente; licencias asociadas en varios casos a cuentas por cargo.','Trazabilidad, baja y MFA dependen de una cuenta que no representa claramente a una persona.'),
        ('Correo','Zimbra con direcciones como cargo@midominio.com.co.','Puede conservarse; correo e identidad no tienen que ser la misma cuenta.'),
        ('Portátiles','Inicio con cuenta local; luego FortiClient VPN y RDP al PC de oficina.','Dos inicios de sesión, menor control del equipo y recuperación compleja fuera de la oficina.'),
        ('Autenticación','Presencia de SMS/voz y contraseñas.','Exposición a phishing, SIM swap y cambios regulatorios/tecnológicos.'),
    ],[1900,3400,4060])
    doc.add_heading('3. Por qué es importante',1)
    table(doc,['Objetivo','Valor para la organización'],[
        ('Reducir incidentes','Passkeys y Windows Hello reducen el valor de robar o reutilizar una contraseña.'),
        ('Responsabilidad individual','Cada acceso, licencia y método de recuperación queda asociado a una persona.'),
        ('Continuidad','La autenticación de Microsoft 365 no depende de que los DC estén disponibles para cada acceso.'),
        ('Trabajo remoto','El portátil se convierte en equipo corporativo administrado, no en un equipo aislado que solo ejecuta una VPN.'),
        ('Cumplimiento','Se eliminan SMS y voz antes del retiro obligatorio y se documentan controles compensatorios.'),
        ('Experiencia','Un usuario personal, Windows Hello y SSO disminuyen solicitudes repetidas de credenciales.'),
    ],[2200,7160])
    callout(doc,'Riesgo de no actuar','Después del 1 de febrero de 2027, un usuario cuyo único MFA sea SMS o voz quedará bloqueado hasta registrar una passkey. La organización además conservaría cuentas compartidas, trazabilidad limitada y portátiles no gobernados.',fill='FDECEC',color=RED)
    doc.add_heading('4. Solución propuesta',1)
    doc.add_paragraph('Se conservarán MIDOMINIO.LOCAL, Zimbra y FortiGate. El cambio consiste en añadir una capa coherente de identidad y control:')
    numbered(doc,[
        'Agregar midominio.com.co como sufijo UPN, sin renombrar el dominio interno.',
        'Asignar a cada persona una identidad nombre.apellido@midominio.com.co y una licencia individual cuando corresponda.',
        'Sincronizar identidades desde AD hacia Entra mediante Password Hash Synchronization.',
        'Registrar los equipos corporativos mediante Microsoft Entra Hybrid Join.',
        'Adoptar Windows Hello for Business en los equipos y passkeys en Microsoft 365.',
        'Integrar FortiGate/FortiClient con Entra mediante SAML y MFA, sujeto a compatibilidad.',
        'Extender administración, cifrado, actualizaciones y protección de endpoint a portátiles.',
        'Sustituir la caducidad periódica de contraseña por una política moderna con controles compensatorios.'
    ])
    doc.add_heading('5. Qué cambia al iniciar sesión',1)
    table(doc,['Escenario','Hoy','Objetivo'],[
        ('PC de oficina','MIDOMINIO\\nombre.apellido + contraseña.','nombre.apellido@midominio.com.co con Windows Hello; contraseña disponible para contingencia.'),
        ('Microsoft 365','Cuenta de cargo y SMS/MFA según configuración.','Identidad personal + passkey; SSO en equipo corporativo.'),
        ('Portátil remoto','Cuenta local → FortiClient → credencial VPN → RDP.','Cuenta corporativa/Hello → VPN corporativa con identidad y MFA → RDP, manteniendo el modelo de acceso remoto inicialmente.'),
        ('Recuperación','Cambio manual y métodos dependientes del teléfono.','Procedimiento verificado con Temporary Access Pass, llave de respaldo y mesa de ayuda.'),
    ],[1800,3500,4060])
    doc.add_paragraph('La fase inicial mantiene el Escritorio remoto al PC de oficina. No se abrirá RDP directamente a Internet. Una evolución posterior podrá evaluar acceso directo a aplicaciones, ZTNA o retiro de dependencias innecesarias de RDP.')
    doc.add_heading('6. Extensión del entorno corporativo a portátiles',1)
    doc.add_paragraph('El objetivo no es solamente conectar una red: es llevar identidad, configuración y protección al dispositivo remoto.')
    bullets(doc,[
        'Equipo unido al dominio y registrado de forma híbrida, o administrado por Entra/Intune según el modelo aprobado.',
        'BitLocker con clave de recuperación custodiada por TI, TPM y Secure Boot.',
        'Windows Hello, EDR/antivirus, firewall, parches y eliminación de administradores locales innecesarios.',
        'FortiClient administrado por EMS cuando el licenciamiento y la arquitectura lo permitan.',
        'VPN antes del inicio de sesión para contactar el dominio en escenarios de aprovisionamiento o cambio de contraseña.',
        'Preferencia por un túnel previo basado en certificado de equipo y autenticación posterior del usuario con Entra/MFA.',
        'Acceso RDP limitado al equipo asignado, con reglas de firewall, registros y segmentación.'
    ])
    callout(doc,'Decisión de arquitectura VPN','Fortinet documenta SAML con Microsoft Entra y VPN antes del inicio de Windows, incluso con certificados de equipo mediante EMS. La opción final requiere conocer modelo de FortiGate, versión de FortiOS, modalidad SSL/IPsec, versión/licencia de FortiClient, existencia de EMS y PKI. No se comprometerá su disponibilidad sin una prueba de concepto.')
    doc.add_heading('7. Política moderna de contraseñas',1)
    doc.add_paragraph('Microsoft Research explica que la caducidad frecuente lleva a patrones previsibles: usuarios incrementan números, cambian una letra o reutilizan estructuras. Microsoft 365 desaconseja la expiración periódica para cuentas cloud; NIST SP 800-63B-4 también indica que no deben exigirse cambios periódicos sin evidencia de compromiso.')
    table(doc,['Control propuesto','Política'],[
        ('Longitud','Mínimo objetivo de 15 caracteres; se favorecen frases de paso memorables y únicas.'),
        ('Caducidad','Sin vencimiento periódico para usuarios, una vez desplegados los controles compensatorios.'),
        ('Cambio obligatorio','Inmediato ante sospecha, filtración, restablecimiento, malware o instrucción de TI.'),
        ('Bloqueo de débiles','Microsoft Entra Password Protection para impedir términos comunes, comprometidos y propios de la organización.'),
        ('MFA','Passkey/Windows Hello; SMS y voz se retiran.'),
        ('Reutilización','Prohibida fuera de la organización; historial y monitoreo internos.'),
        ('Administradores','Llaves FIDO2 y cuentas separadas; política más estricta según riesgo.'),
        ('Cuentas de servicio','Contraseñas administradas/gMSA o certificados; no se aplica la misma política humana.'),
    ],[2600,6760])
    callout(doc,'Condición de seguridad','No se recomienda configurar “la contraseña nunca expira” antes de activar passkeys/MFA, bloqueo de contraseñas débiles, monitoreo, procedimiento de respuesta y protección de cuentas privilegiadas. La mejora es un paquete de controles, no un ajuste aislado.',fill='FFF4D6',color=GOLD)
    doc.add_heading('8. Costos y recursos',1)
    doc.add_paragraph('Las cifras siguientes son escenarios de planeación, no cotizaciones. Valores públicos de referencia en USD, compromiso anual, antes de impuestos y conversión a COP.')
    table(doc,['Concepto','Referencia','Estimación para 40 usuarios'],[
        ('Entra ID P1 independiente','USD 7 usuario/mes.','USD 280/mes; USD 3.360/año.'),
        ('Business Premium','USD 22 usuario/mes con Teams.','USD 880/mes; USD 10.560/año si se licencia a los 40.'),
        ('Incremento Standard → Premium','Precios públicos: USD 14 → USD 22.','USD 8 usuario/mes; USD 320/mes si los 40 ya tuvieran Standard.'),
        ('Llaves FIDO2','Presupuesto orientativo, sujeto a proveedor.','Dos por administrador y para usuarios sin dispositivo compatible; cantidad tras inventario.'),
        ('FortiClient EMS / Fortinet','Depende del contrato, versión y número de endpoints.','Cotización obligatoria con partner; podría ser costo nuevo.'),
        ('Servidor Entra Connect','VM miembro dedicada recomendada.','Sin compra de hardware si existe capacidad y licenciamiento Windows.'),
        ('Servicios profesionales','Diseño, piloto, despliegue y documentación.','80–140 horas como rango inicial; ajustar tras descubrimiento.'),
        ('Trabajo interno','TI, comunicaciones y soporte a usuarios.','0,25–0,5 FTE durante 4–6 meses, con mayor carga en pilotos.'),
    ],[2200,3000,4160],9)
    doc.add_paragraph('Modelo recomendado para evaluar: Business Premium para quienes requieren Office y administración completa del portátil; Entra ID P1 para usuarios interactivos sin Office que queden sujetos a Acceso Condicional. La asignación exacta exige inventario y validación contractual.')
    doc.add_heading('9. Tiempo y etapas',1)
    table(doc,['Etapa','Duración','Resultado'],[
        ('Descubrimiento y decisión','2 semanas','Inventario, licencias, Fortinet, riesgos y diseño aprobado.'),
        ('Identidad y passkeys piloto','2–3 semanas','Cuentas personales, administradores protegidos y usuarios piloto.'),
        ('Entra Connect y dispositivos','3–5 semanas','Sincronización, SSO e Hybrid Join validados.'),
        ('Portátiles y VPN','3–6 semanas','Piloto de equipo corporativo remoto y SAML/prelogon según compatibilidad.'),
        ('Despliegue general','4–6 semanas','Migración por grupos, formación y soporte.'),
        ('Estabilización','2–4 semanas','Auditoría, eliminación de SMS/voz y cierre.'),
    ],[2700,1600,5060])
    doc.add_paragraph('Duración total estimada: 4 a 6 meses. La línea de passkeys debe comenzar inmediatamente y concluir antes de la fecha interna propuesta del 15 de diciembre de 2026.')
    doc.add_heading('10. Riesgos y mitigaciones',1)
    table(doc,['Riesgo','Mitigación'],[
        ('Duplicar o asociar incorrectamente cuentas','Tabla de correspondencia, OU piloto y Entra Connect en staging.'),
        ('Bloquear usuarios al exigir passkey','Piloto, TAP, método de respaldo y soporte reforzado.'),
        ('Incompatibilidad Fortinet','Inventario de versiones, matriz de soporte y prueba de concepto.'),
        ('Cambiar contraseñas sin controles compensatorios','No modificar caducidad hasta cumplir el criterio técnico de salida.'),
        ('Costos superiores','Licenciamiento por perfiles, cotización local y aprobación por etapa.'),
        ('Resistencia del usuario','Comunicación simple, demostraciones y acompañamiento.'),
    ],[3300,6060])
    doc.add_heading('11. Indicadores de éxito',1)
    bullets(doc,[
        '100 % de personas con identidad individual y propietario documentado.',
        '100 % de administradores con cuenta separada y FIDO2.',
        '0 usuarios dependientes de SMS o voz antes del 15 de diciembre de 2026.',
        '100 % de portátiles corporativos cifrados, inventariados y bajo política.',
        '100 % de accesos VPN asociados a una persona y protegidos con MFA compatible.',
        'Reducción de cuentas locales y compartidas a excepciones justificadas.',
        'Procedimientos de alta, cambio, retiro y recuperación probados.'
    ])
    doc.add_heading('12. Aprobación solicitada',1)
    callout(doc,'Recomendación','Aprobar la fase de descubrimiento y piloto, autorizar cotizaciones de Microsoft/Fortinet/llaves FIDO2 y designar un patrocinador ejecutivo. La aprobación del despliegue total se realiza al superar los criterios del piloto y confirmar el costo final.',fill='E8F4EF',color=GREEN)
    sources(doc)
    path=OUT/'Proyecto_Identidad_Acceso_Seguro_Ejecutivo.docx'; doc.save(path); return path

def technical():
    doc=setup(Document(),'Proyecto de identidad y acceso seguro | Plan técnico')
    title(doc,'Plan técnico de modernización de identidad y acceso','Diseño, dependencias, ejecución, validación y reversión','Documento técnico')
    callout(doc,'Objetivo técnico','Integrar MIDOMINIO.LOCAL con Microsoft Entra, conservar Zimbra, normalizar identidades personales, desplegar passkeys y Windows Hello, gobernar portátiles y elevar FortiClient VPN con identidad/MFA sin exponer RDP a Internet.')
    doc.add_heading('1. Alcance y principios',1)
    bullets(doc,[
        'No renombrar el bosque/dominio MIDOMINIO.LOCAL.',
        'No migrar Zimbra ni cambiar MX como parte de este proyecto.',
        'No exponer TCP/3389 directamente a Internet.',
        'No realizar emparejamientos masivos sin piloto y staging.',
        'No eliminar la caducidad de contraseñas antes de los controles compensatorios.',
        'No asumir compatibilidad Fortinet sin registrar modelo, versiones y licencias.',
        'Aplicar mínimo privilegio, cambios por grupos y reversión documentada.'
    ])
    doc.add_heading('2. Arquitectura lógica objetivo',1)
    doc.add_paragraph('Identidad primaria: nombre.apellido@midominio.com.co. El sAMAccountName MIDOMINIO\\nombre.apellido permanece. Zimbra conserva cargo@midominio.com.co como buzón funcional. La licencia de Microsoft 365 y los métodos MFA se asignan a la persona.')
    table(doc,['Plano','Componentes','Flujo'],[
        ('Identidad','AD DS, Entra Connect, Entra ID','AD es origen para usuarios sincronizados; PHS suministra autenticación cloud resiliente.'),
        ('Dispositivo','AD Join + Entra Hybrid Join + gestión','El dispositivo obtiene identidad cloud y aplica cifrado, cumplimiento y protección.'),
        ('Autenticación','Windows Hello, passkeys, FIDO2, TAP','Credenciales resistentes al phishing; contraseña queda como contingencia durante transición.'),
        ('Acceso remoto','FortiClient, FortiGate, Entra SAML, EMS/PKI si aplica','Equipo confiable establece conectividad; usuario se autentica individualmente y accede solo a recursos autorizados.'),
        ('Aplicaciones','RDP y recursos AD','RDP se mantiene detrás de VPN; Kerberos y grupos AD conservan autorización.'),
    ],[1500,3100,4760])
    doc.add_heading('3. Levantamiento obligatorio',1)
    doc.add_heading('3.1 Active Directory',2)
    bullets(doc,[
        'Bosque, dominio, niveles funcionales, sitios, subredes, FSMO, catálogo global y DNS.',
        'Salud: dcdiag /e /c, repadmin /replsummary, repadmin /showrepl * y eventos.',
        'UPN, mail, proxyAddresses, duplicados, cuentas inactivas, administradores y servicios.',
        'OU, GPO, delegaciones, grupos privilegiados, LAPS, PKI y respaldos de estado del sistema.',
        'Windows 10/11 soportado, TPM, Secure Boot, BitLocker y propiedad de cada equipo.'
    ])
    doc.add_heading('3.2 Microsoft 365 y Entra',2)
    bullets(doc,[
        'Dominio verificado, usuarios cloud-only, UPN, Object ID, licencias, roles y métodos de autenticación.',
        'Mapa persona ↔ cuenta AD ↔ cuenta M365 ↔ buzón Zimbra ↔ dispositivos.',
        'Cuentas compartidas, huérfanas, de servicio, invitadas y eliminadas recuperables.',
        'Edición Entra actual, licencias por usuario, Secure Score y registros de inicio.'
    ])
    doc.add_heading('3.3 Fortinet y conectividad',2)
    table(doc,['Dato','Valor a registrar'],[
        ('FortiGate','Modelo, FortiOS, HA, soporte, firmware recomendado y fecha de fin de soporte.'),
        ('VPN','SSL o IPsec, puertos, split/full tunnel, grupos, LDAP/RADIUS, portales y concurrencia.'),
        ('FortiClient','Versión, gratuito/licenciado, XML, distribución, SAML y soporte prelogon.'),
        ('EMS','Existencia, versión, licencia, ubicación, certificados y acceso desde Internet.'),
        ('PKI','AD CS existente, plantillas, autoenrollment, CRL/OCSP y renovación.'),
        ('RDP','Reglas, NLA, grupos permitidos, origen/destino, logs y restricción equipo-asignado.'),
    ],[2200,7160])
    doc.add_heading('4. Diseño de identidad',1)
    doc.add_heading('4.1 Convenciones',2)
    table(doc,['Objeto','Convención'],[
        ('Usuario','nombre.apellido@midominio.com.co'),('sAMAccountName','nombre.apellido'),('Administrador cloud','adm.nombre.apellido@midominio.com.co'),('Emergencia','Dos identidades cloud-only de nombres discretos'),('Servicio local','svc-aplicacion o gMSA; sin uso interactivo'),('Buzón funcional','cargo@midominio.com.co en Zimbra')
    ],[2600,6760])
    doc.add_heading('4.2 Sufijo UPN',2)
    doc.add_paragraph('Agregar midominio.com.co en Active Directory Domains and Trusts o mediante:')
    p=doc.add_paragraph(); p.style='Intense Quote'; p.add_run('Get-ADForest | Select-Object UPNSuffixes\nSet-ADForest -Identity "MIDOMINIO.LOCAL" -UPNSuffixes @{Add="midominio.com.co"}')
    doc.add_paragraph('Actualizar inicialmente solo el piloto:')
    p=doc.add_paragraph(); p.style='Intense Quote'; p.add_run('Set-ADUser -Identity "juan.perez" -UserPrincipalName "juan.perez@midominio.com.co"')
    doc.add_heading('4.3 Reglas de emparejamiento',2)
    bullets(doc,[
        'Caso 1: cuenta de cargo usada exclusivamente por una persona: evaluar cambio de UPN cloud antes de sincronizar para preservar objeto/licencia.',
        'Caso 2: cuenta compartida: crear identidades personales; retirar licencia y sesiones de la compartida; Zimbra puede conservar el buzón.',
        'Caso 3: administrador: separar cuenta normal y privilegiada.',
        'Caso 4: servicio: excluir del flujo humano y migrar a gMSA/certificado cuando sea posible.',
        'No asignar la misma proxyAddress a varios objetos. No forzar ImmutableId sin evidencia y aprobación.'
    ])
    doc.add_heading('5. Entra Connect Sync',1)
    doc.add_heading('5.1 Servidor y método',2)
    bullets(doc,[
        'VM/servidor miembro dedicado, parchado, acceso administrativo restringido y respaldo de configuración.',
        'Instalación personalizada; Password Hash Synchronization; Seamless SSO sujeto a prueba.',
        'Filtrado inicial por OU; cuenta de emergencia y administradores cloud fuera del alcance.',
        'Staging mode durante revisión; no exportar hasta validar add/update/delete y join.'
    ])
    doc.add_heading('5.2 Secuencia',2)
    numbered(doc,[
        'Verificar midominio.com.co en Entra con TXT sin cambiar MX de Zimbra.',
        'Preparar OU Usuarios-Piloto y Equipos-Piloto con objetos reales autorizados.',
        'Instalar Entra Connect en modo personalizado y staging.',
        'Revisar Synchronization Service Manager: conectores, metaverse, coincidencias y eliminaciones.',
        'Exportar la configuración y registrar versión/cuenta/servidor.',
        'Desactivar staging para el piloto; ejecutar sincronización.',
        'Validar inicio cloud, cambio de contraseña local, bloqueo y baja.',
        'Ampliar OU por lotes y observar registros.'
    ])
    doc.add_heading('5.3 Pruebas y reversión',2)
    table(doc,['Prueba','Resultado esperado','Reversión'],[
        ('Usuario existente','Un solo objeto, licencia preservada y origen AD.','Volver a staging; no borrar; corregir atributos.'),
        ('Contraseña','Cambio AD funciona en cloud después de sincronizar.','Revisar PHS/agente; mantener acceso de emergencia.'),
        ('Baja','Deshabilitar AD impide nuevo acceso y se revocan sesiones.','Rehabilitación documentada si fue prueba.'),
        ('Zimbra','Correo sigue llegando a cargo@.','Revertir solo atributo/UPN afectado; no tocar MX.'),
    ],[2200,3700,3460])
    doc.add_heading('6. Passkeys y administración privilegiada',1)
    numbered(doc,[
        'Crear grupos PASSKEY-PILOTO, PASSKEY-PRODUCCION, ADMIN y exclusión temporal.',
        'Crear dos cuentas de emergencia cloud-only, monitoreadas y probadas.',
        'Habilitar Temporary Access Pass para incorporación y recuperación.',
        'Habilitar Passkey (FIDO2), autoservicio y perfiles permitidos para el piloto.',
        'Entregar dos llaves FIDO2 a cada administrador y separar cuenta administrativa.',
        'Registrar 5–7 usuarios piloto, probar navegadores, Office y recuperación.',
        'Ejecutar campaña de registro; medir adopción.',
        'Retirar SMS/voz por grupos después de confirmar dos rutas de recuperación.'
    ])
    doc.add_heading('7. Hybrid Join y Windows Hello',1)
    doc.add_heading('7.1 Microsoft Entra Hybrid Join',2)
    numbered(doc,[
        'Configurar Device Options en Entra Connect y el Service Connection Point.',
        'Sincronizar exclusivamente la OU Equipos-Piloto.',
        'Esperar la tarea automática de registro o reiniciar durante ventana aprobada.',
        'Ejecutar dsregcmd /status como usuario y validar DomainJoined=YES, AzureAdJoined=YES y AzureAdPrt=YES.',
        'Comparar DeviceId con el objeto visible en Entra y revisar equipos pendientes/duplicados.'
    ])
    doc.add_heading('7.2 Windows Hello for Business',2)
    bullets(doc,[
        'Modelo: Hybrid + Cloud Kerberos Trust; no mezclar con certificate trust.',
        'Crear objeto Microsoft Entra Kerberos mediante procedimiento oficial y cuenta privilegiada temporal.',
        'GPO piloto: Use Windows Hello for Business, Use cloud trust for on-premises authentication, TPM y biometría según política.',
        'PIN mínimo de 6 para desbloqueo local; el secreto permanece ligado al equipo. La contraseña corporativa sigue la política de 15 caracteres.',
        'Validar acceso Kerberos a archivos, impresión, aplicaciones y RDP.'
    ])
    doc.add_heading('8. Portátiles y FortiClient',1)
    doc.add_heading('8.1 Estado objetivo del endpoint',2)
    table(doc,['Control','Criterio'],[
        ('Identidad','AD Join + Hybrid Join, o Entra Join/Intune si se aprueba una arquitectura cloud-managed.'),
        ('Inicio','Cuenta corporativa + Windows Hello; la cuenta local queda solo como soporte controlado y administrada por LAPS.'),
        ('Cifrado','BitLocker activo; clave en repositorio corporativo; TPM/Secure Boot.'),
        ('Protección','EDR, firewall, parches, navegador empresarial y sin admin local permanente.'),
        ('VPN','Perfil gestionado; certificado de servidor válido; MFA; logs; sin guardar contraseña.'),
        ('RDP','NLA, grupo individual, destino asignado y acceso solo por VPN.'),
    ],[2300,7060])
    doc.add_heading('8.2 Flujo recomendado en dos etapas',2)
    numbered(doc,[
        'Prelogon técnico: el equipo establece un túnel limitado mediante certificado de máquina para contactar DC/DNS/gestión. Requiere EMS y PKI compatibles.',
        'Inicio de Windows: el usuario entra con su identidad corporativa y Windows Hello.',
        'Autenticación de usuario VPN: FortiGate usa Microsoft Entra como IdP SAML y aplica MFA/Acceso Condicional cuando la combinación sea compatible.',
        'Acceso de trabajo: el usuario abre RDP únicamente a su PC asignado; no obtiene acceso lateral general.',
        'Fin de sesión: se registran eventos, se desconecta el túnel según política y el portátil permanece protegido fuera de red.'
    ])
    callout(doc,'Advertencia técnica','SAML interactivo y túnel automático previo al logon no siempre son el mismo flujo. Para conectividad desatendida antes del usuario, Fortinet documenta certificado de equipo. Después se puede establecer o autorizar el contexto del usuario. La PoC debe demostrar la transición sin crear un túnel excesivamente privilegiado.',fill='FFF4D6',color=GOLD)
    doc.add_heading('8.3 Matriz de prueba Fortinet',2)
    table(doc,['Caso','Validación'],[
        ('Internet doméstico','DNS, certificado público, SAML, MFA y conexión estable.'),
        ('Contraseña cambiada','Prelogon alcanza DC y actualiza credencial sin cuenta local.'),
        ('Equipo robado','BitLocker bloquea datos; certificado/dispositivo/tokens se revocan.'),
        ('Usuario retirado','Cuenta deshabilitada y sesiones revocadas impiden VPN/RDP.'),
        ('Sin Internet','Windows Hello/caché permite acceso local; recursos corporativos no están disponibles.'),
        ('Caída de Entra','Procedimiento de continuidad definido; no se amplía acceso anónimo.'),
        ('Certificado expirado','Renovación previa y alerta; proceso de recuperación controlado.'),
    ],[2400,6960])
    doc.add_heading('9. Acceso Condicional',1)
    table(doc,['Política','Modo inicial','Exclusiones/condiciones'],[
        ('Bloquear autenticación heredada','Report-only','Emergencia temporal; validar aplicaciones.'),
        ('MFA para recursos cloud','Report-only → piloto','Dos emergencias; grupos de servicio justificados.'),
        ('Resistente al phishing para administradores','Piloto → On','Solo FIDO2/Hello/passkeys compatibles.'),
        ('Registro de información de seguridad','Piloto','TAP y ubicación/dispositivo autorizado.'),
        ('Dispositivo conforme','Fase posterior','Requiere gestión/compliance; no confundir Hybrid Join con conformidad.'),
        ('FortiGate Enterprise App','Piloto','Usuarios VPN, contexto de dispositivo según compatibilidad SAML.'),
    ],[2900,1800,4660])
    doc.add_heading('10. Política técnica de contraseñas',1)
    doc.add_heading('10.1 Política objetivo',2)
    table(doc,['Parámetro','Objetivo','Condición'],[
        ('Longitud mínima','15 caracteres','Habilitar primero Relax minimum password length limits en sistemas compatibles.'),
        ('Longitud máxima aceptada','Al menos 64 cuando la aplicación lo permita','Validar sistemas heredados antes de prometer compatibilidad.'),
        ('Edad máxima','0 / sin caducidad periódica','Solo después del gate de controles compensatorios.'),
        ('Edad mínima','1 día','Evita recorrer rápidamente el historial.'),
        ('Historial','24 contraseñas','Complementario; no sustituye bloqueo de comprometidas.'),
        ('Composición','Evitar reglas arbitrarias a largo plazo','Mantener temporalmente si aún no existe Password Protection.'),
        ('Contraseñas débiles','Bloqueo global + lista personalizada','Desplegar Microsoft Entra Password Protection en AD DS.'),
        ('Bloqueo de cuenta','Umbral inicial 10; duración 15 min','Validar con aplicaciones y mesa de ayuda.'),
        ('Cambio forzado','Compromiso real, restablecimiento o evento de riesgo','Revocar tokens y revisar endpoint simultáneamente.'),
    ],[2200,2500,4660],9)
    doc.add_heading('10.2 Gate antes de retirar expiración',2)
    bullets(doc,[
        'Passkey/MFA desplegado y probado en la población afectada.',
        'Administradores con cuentas separadas y FIDO2.',
        'Microsoft Entra Password Protection funcionando en todos los DC, con agentes redundantes y registros revisados.',
        'Procedimiento de detección y respuesta a credencial comprometida.',
        'Bloqueo de autenticación heredada y monitoreo de inicios.',
        'Inventario de servicios y tareas programadas para evitar contraseñas humanas embebidas.',
        'Aprobación de cumplimiento, auditoría y aseguradora si existen requisitos contractuales de caducidad.'
    ])
    doc.add_heading('10.3 Despliegue de Password Protection',2)
    numbered(doc,[
        'Confirmar licenciamiento Entra requerido para usuarios beneficiados por la protección on-premises.',
        'Instalar al menos dos proxy agents en servidores miembro y registrar el bosque.',
        'Instalar el DC agent en los tres DC, escalonando reinicios si fueran requeridos.',
        'Configurar lista prohibida personalizada: razón social, marcas, ciudades, productos y patrones previsibles, sin incluir datos sensibles en el documento.',
        'Iniciar en Audit mode; observar eventos y ajustar.',
        'Pasar a Enforced; medir rechazos y soporte.',
        'Solo entonces aplicar la política de no caducidad a usuarios y separar políticas de servicios/administradores.'
    ])
    doc.add_heading('11. Licenciamiento y presupuesto técnico',1)
    table(doc,['Perfil','Opción recomendada','Cantidad a confirmar'],[
        ('Usuario con Office + portátil administrado','Microsoft 365 Business Premium.','Número de usuarios reales, no número de cuentas por cargo.'),
        ('Usuario AD sin Office pero sujeto a CA','Microsoft Entra ID P1.','Usuarios interactivos alcanzados por la política.'),
        ('Administrador','Licencia correspondiente + dos FIDO2.','Número de personas, no número de roles.'),
        ('Portátil','Gestión/EDR compatible + FortiClient/EMS si aplica.','Número de endpoints remotos.'),
        ('Entra Connect','Windows Server/VM y operación.','Una instancia activa; estrategia de staging/recuperación.'),
        ('Fortinet','EMS, VPN/ZTNA y soporte según contrato.','Cotización del partner con matriz de versiones.'),
    ],[2100,4600,2660])
    doc.add_paragraph('Referencias públicas: Entra ID P1 USD 7 usuario/mes; Business Premium USD 22 usuario/mes con Teams y compromiso anual. Para 40 usuarios: USD 3.360/año de P1 o USD 10.560/año de Premium. Estos totales no representan necesariamente costo incremental: se deben descontar licencias actuales y evitar duplicidades.')
    doc.add_heading('12. Plan de ejecución',1)
    table(doc,['Fase','Semanas','Entregables técnicos','Gate'],[
        ('0 Descubrimiento','1–2','Inventarios, salud AD, licencias, Fortinet y mapa de cuentas.','Datos completos; backup válido.'),
        ('1 Seguridad tenant','2–3','Emergencia, admins FIDO2, TAP y passkey piloto.','Recuperación probada.'),
        ('2 Identidad','2–4','UPN, coincidencias y normalización de cuentas.','Sin duplicados.'),
        ('3 Sincronización','3–6','Entra Connect staging/piloto/producción.','Ciclo de vida probado.'),
        ('4 Dispositivos','5–9','Hybrid Join y Windows Hello piloto.','Kerberos/SSO válidos.'),
        ('5 Fortinet','6–12','PoC SAML, MFA, prelogon y RDP restringido.','Matriz de pruebas aprobada.'),
        ('6 Políticas','9–16','CA report-only, Password Protection y GPO.','Telemetría sin impacto crítico.'),
        ('7 Producción','13–20','Despliegue, retiro SMS/voz y documentación.','KPIs y aceptación.'),
    ],[1500,950,4660,2250],8.5)
    doc.add_heading('13. RACI resumido',1)
    table(doc,['Actividad','Responsable','Aprueba','Consulta'],[
        ('Arquitectura y cambios AD','Administrador de infraestructura','Líder TI','Proveedor/Microsoft'),
        ('FortiGate/FortiClient','Administrador de seguridad/red','Líder TI','Partner Fortinet'),
        ('Licenciamiento','Compras + TI','Dirección','Partner Microsoft'),
        ('Política de contraseñas','Seguridad/TI','Dirección','Auditoría/Legal/RR. HH.'),
        ('Comunicación y adopción','TI + RR. HH.','Patrocinador','Jefes de área'),
        ('Pruebas de usuario','Usuarios piloto','Líder TI','Mesa de ayuda'),
    ],[2700,2500,1800,2360])
    doc.add_heading('14. Evidencias de aceptación',1)
    bullets(doc,[
        'Exportes antes/después de usuarios, licencias, roles y métodos.',
        'Reporte de Entra Connect sin errores y configuración respaldada.',
        'dsregcmd /status de cada equipo piloto con estados esperados.',
        'Capturas/logs de passkey, Hello, Kerberos, SAML VPN y RDP autorizado.',
        'Prueba de baja: revocación de Entra, VPN y RDP.',
        'Prueba de recuperación: TAP/llave de respaldo sin degradar a SMS.',
        'Reporte de BitLocker/EDR/parches de portátiles.',
        'Eventos de Password Protection en auditoría y cumplimiento.',
        'Plan de reversión firmado por cambio.'
    ])
    doc.add_heading('15. Datos pendientes para cerrar el diseño',1)
    table(doc,['Dato pendiente','Impacto'],[
        ('Modelo de FortiGate y versión FortiOS','Compatibilidad SAML, SSL/IPsec y soporte.'),
        ('Versión/licencia FortiClient y existencia de EMS','Prelogon, administración, certificados y costos.'),
        ('Tipo exacto de licencias Microsoft','Costo incremental y derechos de Intune/CA/Office.'),
        ('Cantidad de portátiles y si son compartidos','Licencias, perfiles y llaves FIDO2.'),
        ('Existencia de AD CS/PKI','Prelogon por certificado y operación de ciclo de vida.'),
        ('Requisitos de auditoría/seguro','Posibles excepciones a no caducidad.'),
        ('Versiones Windows y TPM','Compatibilidad Hello y cifrado.'),
        ('Uso de SSL VPN o IPsec','Diseño y versión mínima de FortiClient.'),
    ],[3600,5760])
    sources(doc)
    path=OUT/'Proyecto_Identidad_Acceso_Seguro_Tecnico.docx'; doc.save(path); return path

if __name__=='__main__':
    print(executive().resolve())
    print(technical().resolve())
